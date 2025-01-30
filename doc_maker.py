# workflow:

#1: open all files in project folder

#2: covert files to raw text
#3: combine all text into one long string

#4: chunk text
#5: ai embed/vectorize text chunks
#6: save to db (pandas for temporary saving)

#7: open template document
#8: pull prompts out of document

#9: answer each prompt with RAG
#10: replace prompt with HTML RAG answer

#11: save doc as docx with file name and time stamp


print("Starting doc_maker.py")

#python libraries:
################################################################################
################################################################################

print("Loading standard python libraries")
#standard python libraries:
import os
import re
from datetime import datetime
from typing import Set, Dict, List

print("Loading pip python libraries")
#pip python libraries:
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.shared import qn


print("Loading custom python libraries")
#custom libraries:
from read_doc import read_doc
from ai import embed, chunk_embed_df, RAG_LLM

print("All python libraries have been loaded")
################################################################################
################################################################################


#program variables
################################################################################
################################################################################
input_docs_folder_path = "input_docs"
output_template_folder_path = "output_template"
template_file = "/Users/jakegrim/Desktop/template7.docx"

################################################################################
################################################################################







#main program:
################################################################################
################################################################################
def main():
    print("Now running main program")

    #1: open all files in project folder
    input_docs_files_list = list_files_in_folder(input_docs_folder_path)

    all_doc_text = ""
    for doc in input_docs_files_list:
        doc_file_path = f"{input_docs_folder_path}/{doc}"
        print(f"Reading: {doc_file_path}")

        #2: covert files to raw text
        doc_text = read_doc(doc_file_path)
        #print(f"Doc Text: \n{doc_text}")

        #3: combine all text into one long string
        all_doc_text = all_doc_text + f" ||| DOCUMENT:{doc} TEXT>>> " + doc_text

    input("wait")
    print(all_doc_text)

    #4: chunk text
    #5: ai embed/vectorize text chunks
    #6: save to db (pandas for temporary saving)
    embeded_df = chunk_embed_df(all_doc_text)
    # embeded_df structure:
    """
    {
        'chunk': chunks,
        'embedding': embeddings
    }
    """

    #7: open template document
    
    
    # Get current date and time
    current_datetime = datetime.now()

    # Format the date and time as y_m_d_h_m_s
    formatted_datetime_stamp = current_datetime.strftime("_%Y_%m_%d_%H_%M_%S")

    output_file = f"{output_template_folder_path}/KVG_proposal_{formatted_datetime_stamp}.docx"

    day_month_year = current_datetime.strftime("%d, %B %Y")

    day_month_year_tag = "{%todays_date_dd_month_year%}"

    doc = replace_text_in_docx(template_file, day_month_year_tag, day_month_year)

    #working_doc.save(output_file)
    

    print(f"day_month_year: {day_month_year}")
    #8: pull prompts out of document
    
    # Find prompts with {! !} delimiters
    short_prompts = find_prompts(doc, "{!", "!}")
    print("\nPrompts with {! !}:")
    for prompt in short_prompts:
        question = prompt.replace("{!","").replace("!}","")
        full_prompt_text = f"Provide a short answer the the following question: {question}"
        print(f"{full_prompt_text}")

        #9: answer each prompt with RAG
        rag_llm_answer = RAG_LLM(prompt, embeded_df)

        #10: replace prompt with HTML RAG answer
        doc = replace_text_in_doc(doc, prompt, rag_llm_answer)

    # Find prompts with {! !} delimiters
    big_prompts = find_prompts(doc, "{$", "$}")
    print("\nPrompts with {$ $}:")
    for prompt in big_prompts:
        prompt_text = prompt.replace("{$","").replace("$}","")
        full_prompt_text = f"Our Company KVG is writing a Government contract proposal. Create the following section of the proposal: {prompt_text}"
        print(f"{full_prompt_text}")
        
        #9: answer each prompt with RAG
        rag_llm_answer = RAG_LLM(prompt, embeded_df)

        #10: replace prompt with HTML RAG answer
        doc = replace_text_in_doc(doc, prompt, rag_llm_answer)


    #11: save doc as docx with file name and time stamp
    doc.save(output_file)

    print(f"Saved file as {output_file}")
    print("Done!")
    







################################################################################
################################################################################
################################################################################
################################################################################
################################################################################
################################################################################
################################################################################






#functions
################################################################################
################################################################################
def list_files_in_folder(folder_path):
    """
    List all files in the specified folder.

    :param folder_path: Path to the folder as a string
    :return: A list of file names in the folder
    """
    try:
        # Use os.listdir to get all entries in the directory
        entries = os.listdir(folder_path)
        
        # Filter out directories, keeping only files
        files = [f for f in entries if os.path.isfile(os.path.join(folder_path, f))]
        
        return files
    except FileNotFoundError:
        print(f"The folder '{folder_path}' does not exist.")
        return []
    except PermissionError:
        print(f"You do not have permission to access '{folder_path}'.")
        return []

def get_output_file_name(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
            # Use regex to find the content between the markers
            match = re.search(r'FILE_NAME\{(.*?)\}END_FILE_NAME', content, re.DOTALL)
            
            if match:
                # Group 1 will contain the content between the markers
                extracted_text = match.group(1)
                return extracted_text.strip()  # Remove leading/trailing whitespace
            else:
                print("No content found between the markers.")
                return None
    
    except FileNotFoundError:
        print(f"The file {file_path} was not found.")
        return None
    except IOError:
        print(f"An error occurred while reading the file {file_path}.")
        return None
    
def replace_text_in_paragraph(paragraph, search_text, replace_text):
    """
    Replace text in a paragraph while handling split runs.
    """
    text = paragraph.text
    if search_text in text:
        # Clear all runs
        p_obj = paragraph._p
        for idx in range(len(paragraph.runs)):
            p_obj.remove(paragraph.runs[0]._r)
        
        # Replace the text and add it back as a single run
        new_text = text.replace(search_text, replace_text)
        paragraph.add_run(new_text)
        return True
    return False

def replace_text_in_section(section, search_text, replace_text):
    """
    Replace text in header and footer sections.
    """
    replacements_made = False
    
    # Handle header
    header = section.header
    if header:
        for paragraph in header.paragraphs:
            if replace_text_in_paragraph(paragraph, search_text, replace_text):
                replacements_made = True
        
        # Handle tables in header
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if replace_text_in_paragraph(paragraph, search_text, replace_text):
                            replacements_made = True
    
    # Handle footer
    footer = section.footer
    if footer:
        for paragraph in footer.paragraphs:
            if replace_text_in_paragraph(paragraph, search_text, replace_text):
                replacements_made = True
        
        # Handle tables in footer
        for table in footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if replace_text_in_paragraph(paragraph, search_text, replace_text):
                            replacements_made = True
    
    return replacements_made

def replace_text_in_core_properties(doc, search_text, replace_text):
    """
    Replace text in document core properties (title, subject, etc.)
    """
    replacements_made = False
    
    # Handle document title
    if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
        if search_text in doc.core_properties.title:
            doc.core_properties.title = doc.core_properties.title.replace(search_text, replace_text)
            replacements_made = True
    
    # Handle document subject
    if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
        if search_text in doc.core_properties.subject:
            doc.core_properties.subject = doc.core_properties.subject.replace(search_text, replace_text)
            replacements_made = True
            
    return replacements_made

def replace_text_in_docx(input_path, search_text, replace_text):

    doc = Document(input_path)
    
    return replace_text_in_doc(doc, search_text, replace_text)

def replace_text_in_doc(doc, search_text, replace_text):
    """
    Replace text in a Word document while preserving all formatting and images.
    Includes headers, footers, and document properties.
    
    Args:
        input_path (str): Path to the input Word document
        search_text (str): Text to search for
        replace_text (str): Text to replace with
    
    Returns:
        Document: Modified document object
    """
    try:
        
        # Track if we made any replacements
        replacements_made = False
        
        # Replace in core properties
        if replace_text_in_core_properties(doc, search_text, replace_text):
            replacements_made = True
        
        # Replace in all sections (headers and footers)
        for section in doc.sections:
            if replace_text_in_section(section, search_text, replace_text):
                replacements_made = True
        
        # Iterate through all paragraphs in main document
        for paragraph in doc.paragraphs:
            if replace_text_in_paragraph(paragraph, search_text, replace_text):
                replacements_made = True
        
        # Iterate through all tables in main document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if replace_text_in_paragraph(paragraph, search_text, replace_text):
                            replacements_made = True
        
        if not replacements_made:
            print("Warning: No replacements were made. Check if the search text exists in the document.")
            
        # return the modified document
        return doc
        
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return None
        
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return None


def find_placeholders_in_paragraph(paragraph_text: str) -> Set[str]:
    """
    Find all placeholders matching {% TEXT %} pattern in a paragraph.
    
    Args:
        paragraph_text (str): Text to search in
        
    Returns:
        Set[str]: Set of unique placeholders found
    """
    pattern = r'{%\s*([^%]*?)\s*%}'
    matches = re.finditer(pattern, paragraph_text)
    return {match.group(0) for match in matches}

def count_placeholders(text: str) -> Dict[str, int]:
    """
    Count occurrences of each placeholder in text.
    
    Args:
        text (str): Text to search in
        
    Returns:
        Dict[str, int]: Dictionary with placeholders and their counts
    """
    pattern = r'{%\s*([^%]*?)\s*%}'
    matches = re.finditer(pattern, text)
    counts = {}
    for match in matches:
        placeholder = match.group(0)
        counts[placeholder] = counts.get(placeholder, 0) + 1
    return counts

def find_placeholders_in_docx(file_path: str, count_occurrences: bool = True) -> Dict[str, int]:
    """
    Find all placeholders in a Word document.
    
    Args:
        file_path (str): Path to the Word document
        count_occurrences (bool): Whether to count occurrences of each placeholder
        
    Returns:
        Dict[str, int]: Dictionary of placeholders and their counts (if counting enabled)
                       or placeholders with count 1 (if counting disabled)
    """
    try:
        doc = Document(file_path)
        all_counts = {}
        
        # Search in paragraphs
        for paragraph in doc.paragraphs:
            if count_occurrences:
                counts = count_placeholders(paragraph.text)
                for placeholder, count in counts.items():
                    all_counts[placeholder] = all_counts.get(placeholder, 0) + count
            else:
                placeholders = find_placeholders_in_paragraph(paragraph.text)
                for placeholder in placeholders:
                    all_counts[placeholder] = 1
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if count_occurrences:
                            counts = count_placeholders(paragraph.text)
                            for placeholder, count in counts.items():
                                all_counts[placeholder] = all_counts.get(placeholder, 0) + count
                        else:
                            placeholders = find_placeholders_in_paragraph(paragraph.text)
                            for placeholder in placeholders:
                                all_counts[placeholder] = 1
        
        return all_counts
        
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return {}


################################################################################



def find_prompts_in_text(text: str, start_delim: str, end_delim: str) -> Set[str]:
    """
    Find all placeholders matching {start_delim}TEXT{end_delim} pattern in text.
    
    Args:
        text (str): Text to search in
        start_delim (str): Starting delimiter (e.g., "{!", "{%", "{$")
        end_delim (str): Ending delimiter (e.g., "!}", "%}", "$}")
        
    Returns:
        Set[str]: Set of unique placeholders found
    """
    # Escape special regex characters in delimiters
    start_pattern = re.escape(start_delim)
    end_pattern = re.escape(end_delim)
    
    # Build pattern that matches any character except the end delimiter
    pattern = f'{start_pattern}(.*?){end_pattern}'
    
    try:
        matches = re.finditer(pattern, text)
        return {match.group(0) for match in matches}
    except re.error as e:
        print(f"Regex error: {e}")
        return set()

def find_prompts(doc, start_delim: str, end_delim: str) -> List[str]:
    """
    Find all prompts with custom delimiters in a Word document.
    Searches through main body, headers, footers, tables, and properties.
    
    Args:
        doc: Word document object
        start_delim (str): Starting delimiter (e.g., "{!", "{%", "{$")
        end_delim (str): Ending delimiter (e.g., "!}", "%}", "$}")
        
    Returns:
        List[str]: List of unique prompts found
    """
    unique_prompts = set()
    
    # Check core properties
    if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
        unique_prompts.update(find_prompts_in_text(doc.core_properties.title, start_delim, end_delim))
    
    if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
        unique_prompts.update(find_prompts_in_text(doc.core_properties.subject, start_delim, end_delim))
    
    # Check headers and footers in all sections
    for section in doc.sections:
        # Check header
        header = section.header
        if header:
            for paragraph in header.paragraphs:
                unique_prompts.update(find_prompts_in_text(paragraph.text, start_delim, end_delim))
            
            # Check tables in header
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            unique_prompts.update(find_prompts_in_text(paragraph.text, start_delim, end_delim))
        
        # Check footer
        footer = section.footer
        if footer:
            for paragraph in footer.paragraphs:
                unique_prompts.update(find_prompts_in_text(paragraph.text, start_delim, end_delim))
            
            # Check tables in footer
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            unique_prompts.update(find_prompts_in_text(paragraph.text, start_delim, end_delim))
    
    # Check main document paragraphs
    for paragraph in doc.paragraphs:
        unique_prompts.update(find_prompts_in_text(paragraph.text, start_delim, end_delim))
    
    # Check main document tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    unique_prompts.update(find_prompts_in_text(paragraph.text, start_delim, end_delim))
    
    # Convert set to sorted list for consistent output
    return sorted(list(unique_prompts))



################################################################################
################################################################################



#starts main (should always be at the bottom)
################################################################################
################################################################################
if __name__ == "__main__":
    main()
################################################################################
################################################################################