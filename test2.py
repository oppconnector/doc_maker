from docx import Document
import os
from copy import deepcopy

def replace_text_in_docx(input_path, output_path, search_text, replace_text):
    """
    Replace text in a Word document while preserving all formatting and images.
    
    Args:
        input_path (str): Path to the input Word document
        output_path (str): Path where the modified document will be saved
        search_text (str): Text to search for
        replace_text (str): Text to replace with
    
    Returns:
        bool: True if successful, False if an error occurred
    """
    try:
        # Load the document
        doc = Document(input_path)
        
        # Iterate through all paragraphs
        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                # We need to preserve the runs (formatting)
                inline = paragraph.runs
                # Loop through runs
                for item in inline:
                    if search_text in item.text:
                        # Replace text while preserving formatting
                        item.text = item.text.replace(search_text, replace_text)
        
        # Iterate through all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if search_text in paragraph.text:
                            inline = paragraph.runs
                            for item in inline:
                                if search_text in item.text:
                                    item.text = item.text.replace(search_text, replace_text)
        
        # Save the modified document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return False

def main():
    # Example usage
    input_file = '/Users/jakegrim/Desktop/template.docx'
    output_file = '/Users/jakegrim/Desktop/output5.docx'
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        return
    
    success = replace_text_in_docx(
        input_file,
        output_file,
        "{%solicitation_number%}",
        "123456"
    )
    
    if success:
        print(f"Successfully created {output_file} with replaced text!")
    else:
        print("Failed to process the document.")

if __name__ == "__main__":
    main()




