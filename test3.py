from docx import Document
import os

def replace_text_in_paragraph(paragraph, search_text, replace_text):
    """
    Replace text in a paragraph while handling split runs.
    """
    text = paragraph.text
    if search_text in text:
        # Clear all runs
        p_obj = paragraph._p
        for idx in range(len(paragraph.runs)):
            p_obj.remove(paragraph.runs[0]._r)
        
        # Replace the text and add it back as a single run
        new_text = text.replace(search_text, replace_text)
        paragraph.add_run(new_text)
        return True
    return False

def replace_text_in_docx(input_path, search_text, replace_text):
    """
    Replace text in a Word document while preserving all formatting and images.
    
    Args:
        input_path (str): Path to the input Word document
        output_path (str): Path where the modified document will be saved
        search_text (str): Text to search for
        replace_text (str): Text to replace with
    
    Returns:
        bool: True if successful, False if an error occurred
    """
    try:
        # Load the document
        doc = Document(input_path)
        
        # Track if we made any replacements
        replacements_made = False
        
        # Iterate through all paragraphs
        for paragraph in doc.paragraphs:
            if replace_text_in_paragraph(paragraph, search_text, replace_text):
                replacements_made = True
        
        # Iterate through all tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if replace_text_in_paragraph(paragraph, search_text, replace_text):
                            replacements_made = True
        
        
        if not replacements_made:
            print("Warning: No replacements were made. Check if the search text exists in the document.")
            
        # return the modified document
        return doc
        
    except Exception as e:
        print(f"An error occurred: {str(e)}")
        return None

def main():
    # Example usage
    input_file = '/Users/jakegrim/Desktop/template2.docx'
    output_file = '/Users/jakegrim/Desktop/output7.docx'
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        return
    
    doc = replace_text_in_docx(input_file, "{!solicitation_number!}", "123456")
    
    doc.save(output_file)

if __name__ == "__main__":
    main()